"""
Génération de la lettre de plainte au format PDF (ReportLab) ou Word (python-docx).
"""

import os
from datetime import datetime

from config import DESKTOP, VERSION


def generer_plainte_pdf(profil, vol, destinataire, commune_survol=None):
    """
    Génère un PDF de plainte sur le Bureau et retourne le chemin du fichier créé.
    Lève RuntimeError si ReportLab n'est pas installé.
    """
    if commune_survol is None:
        commune_survol = profil.get("ville", "")
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_RIGHT, TA_JUSTIFY
    except ImportError:
        raise RuntimeError(
            "Le module reportlab n'est pas installé.\n"
            "Ouvrez un terminal et tapez : pip install reportlab")

    prenom   = profil.get("prenom", "").upper()
    nom      = profil.get("nom", "").upper()
    adresse  = profil.get("adresse", "")
    cp       = profil.get("code_postal", "")
    ville_pl = profil.get("ville", "")

    date_vol  = vol.get("date", "")
    heure_vol = vol.get("heure", "")
    indicatif = vol.get("indicatif", "")
    icao24    = vol.get("icao24", "")

    date_sign = datetime.now().strftime("%d/%m/%Y")

    if indicatif and indicatif != "-":
        ref_vol = indicatif.strip()
    elif icao24:
        ref_vol = icao24.strip()
    else:
        ref_vol = "référence inconnue"

    horodatage  = datetime.now().strftime("%Y%m%d_%H%M%S")
    nom_fichier = f"Plainte_{horodatage}.pdf"
    dossier     = DESKTOP if os.path.isdir(DESKTOP) else os.path.expanduser("~")
    chemin      = os.path.join(dossier, nom_fichier)

    doc = SimpleDocTemplate(chemin, pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=3*cm, rightMargin=2.5*cm)

    s_normal = ParagraphStyle('n', fontName='Helvetica',      fontSize=10, leading=14, spaceAfter=2)
    s_bold   = ParagraphStyle('b', fontName='Helvetica-Bold', fontSize=10, leading=14, spaceAfter=2)
    s_right  = ParagraphStyle('r', fontName='Helvetica',      fontSize=10, leading=14, alignment=TA_RIGHT)
    s_objet  = ParagraphStyle('o', fontName='Helvetica-Bold', fontSize=11, leading=15, spaceAfter=4)
    s_body   = ParagraphStyle('c', fontName='Helvetica',      fontSize=11, leading=18, spaceAfter=4, alignment=TA_JUSTIFY)
    s_small  = ParagraphStyle('s', fontName='Helvetica',      fontSize=8,  leading=11,
                               textColor=colors.grey, alignment=TA_RIGHT)

    bleu  = colors.HexColor('#4472C4')
    story = []

    # Expéditeur
    story.append(Paragraph(f"{prenom} {nom}", s_bold))
    if adresse:
        story.append(Paragraph(adresse, s_normal))
    story.append(Paragraph(f"{cp} {ville_pl}".strip(), s_normal))
    story.append(Spacer(1, 0.8*cm))

    # Destinataire
    dest_nom     = destinataire.get("nom", "")
    dest_adresse = destinataire.get("adresse", "")
    dest_cp      = destinataire.get("cp_ville", "")
    dest_bloc = f"{dest_nom}<br/>{dest_adresse}<br/>{dest_cp}"
    story.append(Paragraph(dest_bloc, s_right))
    story.append(Spacer(1, 0.6*cm))

    # Lieu et date
    story.append(Paragraph(f"{ville_pl}, le {date_sign}", s_right))
    story.append(Spacer(1, 0.6*cm))

    # Objet
    heure_obj = heure_vol[:5] if len(heure_vol) >= 5 else heure_vol
    story.append(Paragraph(
        f"Objet : Plainte pour nuisance aérienne — vol du {date_vol} à {heure_obj}", s_objet))

    story.append(HRFlowable(width="100%", thickness=1.5, color=bleu))
    story.append(Spacer(1, 0.5*cm))

    # Corps
    lignes = [
        f"Je soussigné(e) {prenom} {nom}, demeurant au {adresse}, {cp}, {ville_pl},",
        (f"déclare avoir été gêné(e) par un avion volant à basse altitude"
         f" le {date_vol} à {heure_vol} au-dessus de {commune_survol}."),
        None,
        "Je souhaite que ma plainte soit enregistrée et qu'une réponse circonstanciée me soit adressée.",
        None,
        "Si une infraction était constatée, je souhaite que des sanctions soient prises contre les responsables.",
        None,
        f"Pour information, il semblerait que le vol concerné soit le vol : {ref_vol}",
    ]
    for ligne in lignes:
        if ligne:
            story.append(Paragraph(ligne, s_body))
        else:
            story.append(Spacer(1, 0.25*cm))

    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(
        "Veuillez agréer, Madame, Monsieur, l'expression de mes salutations distinguées.", s_body))
    story.append(Spacer(1, 1.5*cm))

    # Signature
    story.append(Paragraph("Signature :", s_normal))
    story.append(Spacer(1, 1.8*cm))
    story.append(Paragraph("_______________________________", s_normal))
    story.append(Paragraph(f"{prenom} {nom}", s_normal))

    doc.build(story)
    return chemin


def generer_plainte_word(profil, vol, destinataire, commune_survol=None):
    """
    Génère un fichier Word (.docx) de plainte sur le Bureau et retourne le chemin.
    Lève RuntimeError si python-docx n'est pas installé.
    """
    if commune_survol is None:
        commune_survol = profil.get("ville", "")
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "Le module python-docx n'est pas installé.\n"
            "Ouvrez un terminal et tapez : pip install python-docx")

    prenom   = profil.get("prenom", "").upper()
    nom      = profil.get("nom", "").upper()
    adresse  = profil.get("adresse", "")
    cp       = profil.get("code_postal", "")
    ville_pl = profil.get("ville", "")

    date_vol  = vol.get("date", "")
    heure_vol = vol.get("heure", "")
    indicatif = vol.get("indicatif", "")
    icao24    = vol.get("icao24", "")

    date_sign = datetime.now().strftime("%d/%m/%Y")

    if indicatif and indicatif != "-":
        ref_vol = indicatif.strip()
    elif icao24:
        ref_vol = icao24.strip()
    else:
        ref_vol = "référence inconnue"

    horodatage  = datetime.now().strftime("%Y%m%d_%H%M%S")
    nom_fichier = f"Plainte_{horodatage}.docx"
    dossier     = DESKTOP if os.path.isdir(DESKTOP) else os.path.expanduser("~")
    chemin      = os.path.join(dossier, nom_fichier)

    document = Document()

    # Marges
    for section in document.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    def _add(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, color=None):
        p = document.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(2)
        return p

    # Expéditeur
    _add(f"{prenom} {nom}", bold=True)
    if adresse:
        _add(adresse)
    _add(f"{cp} {ville_pl}".strip())
    document.add_paragraph()

    # Destinataire (aligné à droite)
    dest_nom     = destinataire.get("nom", "")
    dest_adresse = destinataire.get("adresse", "")
    dest_cp      = destinataire.get("cp_ville", "")
    for ligne in [dest_nom, dest_adresse, dest_cp]:
        if ligne:
            _add(ligne, align=WD_ALIGN_PARAGRAPH.RIGHT)
    document.add_paragraph()

    # Lieu et date
    _add(f"{ville_pl}, le {date_sign}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    document.add_paragraph()

    # Objet
    heure_obj = heure_vol[:5] if len(heure_vol) >= 5 else heure_vol
    p_objet = _add(
        f"Objet : Plainte pour nuisance aérienne — vol du {date_vol} à {heure_obj}",
        bold=True, size=11)

    # Filet bleu
    p_hr = document.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(6)
    run_hr = p_hr.add_run("─" * 80)
    run_hr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run_hr.font.size = Pt(7)

    # Corps
    lignes = [
        f"Je soussigné(e) {prenom} {nom}, demeurant au {adresse}, {cp}, {ville_pl},",
        (f"déclare avoir été gêné(e) par un avion volant à basse altitude"
         f" le {date_vol} à {heure_vol} au-dessus de {commune_survol}."),
        None,
        "Je souhaite que ma plainte soit enregistrée et qu'une réponse circonstanciée me soit adressée.",
        None,
        "Si une infraction était constatée, je souhaite que des sanctions soient prises contre les responsables.",
        None,
        f"Pour information, il semblerait que le vol concerné soit le vol : {ref_vol}",
    ]
    for ligne in lignes:
        if ligne:
            _add(ligne, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        else:
            document.add_paragraph()

    document.add_paragraph()
    _add("Veuillez agréer, Madame, Monsieur, l'expression de mes salutations distinguées.",
         size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    document.add_paragraph()
    document.add_paragraph()
    _add("Signature :")
    document.add_paragraph()
    document.add_paragraph()
    _add("_______________________________")
    _add(f"{prenom} {nom}")

    document.save(chemin)
    return chemin
